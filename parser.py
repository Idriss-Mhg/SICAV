import re
from docx.oxml.ns import qn

SUPPLEMENT_RE = re.compile(r"^SUPPLEMENT\s+\d+\.", re.IGNORECASE)


def _extract_short_name(full_text):
    """
    'SUPPLEMENT 3. CPR Invest – Global' → 'CPR Invest – Global'
    Returns the part after 'SUPPLEMENT N. ', stripped.
    """
    return SUPPLEMENT_RE.sub("", full_text).strip()


def find_compartments(doc):
    """
    Scans doc.paragraphs for lines matching 'SUPPLEMENT N. ...'
    Returns a list of dicts: {name, short_name, start, end}.
      name       : full text  ('SUPPLEMENT 1. CPR Invest – Silver Age')
      short_name : name only  ('CPR Invest – Silver Age')
    """
    paragraphs = doc.paragraphs
    compartments = []

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if SUPPLEMENT_RE.match(text):
            compartments.append({
                "name":       text,
                "short_name": _extract_short_name(text),
                "start":      i,
                "end":        None,
            })

    for j in range(len(compartments) - 1):
        compartments[j]["end"] = compartments[j + 1]["start"] - 1
    if compartments:
        compartments[-1]["end"] = len(paragraphs) - 1

    return compartments


def _get_run_color(para):
    """
    Returns the explicit w:color val of the first run (e.g. '4472C4'),
    or None if no direct color is set (inherited / auto).
    """
    if not para.runs:
        return None
    rPr = para.runs[0]._element.rPr
    if rPr is None:
        return None
    color_elem = rPr.find(qn("w:color"))
    if color_elem is None:
        return None
    val = color_elem.get(qn("w:val"))
    return None if val in (None, "auto") else val


def _is_in_table(para):
    """True if the paragraph lives inside a table cell rather than the document body."""
    parent = para._element.getparent()
    return parent is not None and parent.tag != qn("w:body")


def _is_title_like(para, anchor_para):
    """
    A paragraph marks a new section boundary if it matches the anchor's
    bold state AND its explicit color AND is not a list/indented item.
    """
    if not para.text.strip():
        return False

    # List items and indented paragraphs are sub-content, never section boundaries.
    pPr = para._element.pPr
    if pPr is not None:
        if pPr.find(qn("w:numPr")) is not None:
            return False
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            left = ind.get(qn("w:left")) or ind.get(qn("w:start")) or "0"
            try:
                if int(left) > 0:
                    return False
            except ValueError:
                pass

    anchor_bold  = bool(anchor_para.runs[0].bold) if anchor_para.runs else False
    para_bold    = bool(para.runs[0].bold)         if para.runs        else False
    anchor_color = _get_run_color(anchor_para)
    para_color   = _get_run_color(para)

    return para_bold and para_bold == anchor_bold and para_color == anchor_color


def find_insert_idx(paragraphs, anchor_idx, comp_end, position):
    """
    Returns the paragraph index AFTER WHICH the clause should be inserted.

    position='apres_titre'   → anchor_idx  (right after the anchor title)
    position='apres_section' → last body paragraph of the anchor's section.

    Algorithm: 'pending boundary'
      - Table-cell paragraphs are ignored entirely (doc.paragraphs includes them).
      - When a title-like paragraph is seen, we record a *potential* boundary but
        keep scanning.  If body content follows, the title-like paragraph was a
        sub-heading, not a true boundary, so we clear the pending flag.
      - If no body content follows the last title-like hit, we return the position
        recorded just before that title (= end of the section content).
    """
    if position != "apres_section":
        return anchor_idx

    anchor_para      = paragraphs[anchor_idx]
    last_content     = anchor_idx
    pending_boundary = None   # last_content at time of last title-like hit

    for i in range(anchor_idx + 1, comp_end + 1):
        para = paragraphs[i]
        if not para.text.strip():
            continue
        if _is_in_table(para):            # skip cell paragraphs
            continue
        if _is_title_like(para, anchor_para):
            pending_boundary = last_content   # tentative section end
        else:
            pending_boundary = None           # content after title → sub-heading
            last_content = i

    return pending_boundary if pending_boundary is not None else last_content


def find_anchor(paragraphs, anchor_text, start, end):
    """
    Searches paragraphs[start..end] for the one that best matches anchor_text.
    Strategy (in order of preference):
      1. Exact match (case-insensitive, stripped)
      2. Paragraph starts with anchor text
      3. Anchor text is a substring of paragraph text
    Returns the paragraph index, or None if not found.
    """
    anchor_norm = anchor_text.lower().strip()

    for i in range(start, end + 1):
        para_norm = paragraphs[i].text.lower().strip()
        if para_norm == anchor_norm:
            return i

    for i in range(start, end + 1):
        para_norm = paragraphs[i].text.lower().strip()
        if para_norm.startswith(anchor_norm):
            return i

    for i in range(start, end + 1):
        para_norm = paragraphs[i].text.lower().strip()
        if anchor_norm in para_norm:
            return i

    return None


def match_compartment(comp_name_excel, compartments):
    """
    Matches the short name from the Excel ('CPR Invest – Silver Age')
    against compartments' short_name extracted from the document.
    Strategy: exact → startswith → substring (all case-insensitive).
    """
    name_norm = comp_name_excel.lower().strip()

    for comp in compartments:
        if comp["short_name"].lower() == name_norm:
            return comp

    for comp in compartments:
        short = comp["short_name"].lower()
        if short.startswith(name_norm) or name_norm.startswith(short):
            return comp

    for comp in compartments:
        short = comp["short_name"].lower()
        if name_norm in short or short in name_norm:
            return comp

    return None
