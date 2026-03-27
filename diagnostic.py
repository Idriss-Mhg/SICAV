"""
diagnostic.py — affiche la structure de paragraphes autour des points d'insertion.

Usage:
    python diagnostic.py data/prospectus.docx data/mapping.xlsx
    python diagnostic.py data/prospectus.docx data/mapping.xlsx "defensive"
    python diagnostic.py data/prospectus.docx data/mapping.xlsx "global equity"
"""
import sys
from docx import Document
from docx.oxml.ns import qn
from mapping import load_mapping
from parser import find_compartments, find_anchor, find_insert_idx, _is_in_table


# ── Body-structure helpers ────────────────────────────────────────────────────

def _build_body_map(doc):
    """
    Returns (body_children, para_to_body_pos, body_pos_to_para).

    body_children     : list of direct w:body child elements (w:p and w:tbl)
    para_to_body_pos  : {para._element: index_in_body_children}
    body_pos_to_para  : {index_in_body_children: para_index_in_doc.paragraphs}
    """
    body = doc.element.body
    body_children = list(body)
    para_to_body_pos = {}
    body_pos_to_para = {}

    para_idx = 0
    for bi, child in enumerate(body_children):
        if child.tag == qn("w:p"):
            para_to_body_pos[child] = bi
            body_pos_to_para[bi] = para_idx
            para_idx += 1
        elif child.tag == qn("w:tbl"):
            # Count all paragraphs inside the table cells
            for tc_p in child.iter(qn("w:p")):
                para_to_body_pos[tc_p] = bi   # all point to the table's body pos
                body_pos_to_para.setdefault(bi, para_idx)  # first para of table
                para_idx += 1

    return body_children, para_to_body_pos, body_pos_to_para


def dump_body_structure(doc, paragraphs, comp_start, comp_end):
    """
    Shows the raw sequence of paragraphs and tables in the document body
    for the given compartment range.  Useful for understanding why insertions
    land in the wrong place relative to tables.
    """
    body = doc.element.body
    body_children = list(body)

    # Build para_element → para_index map (body-level paragraphs only)
    elem_to_idx = {}
    idx = 0
    for child in body_children:
        if child.tag == qn("w:p"):
            elem_to_idx[id(child)] = idx
            idx += 1
        elif child.tag == qn("w:tbl"):
            for tc_p in child.iter(qn("w:p")):
                idx += 1   # count but don't map (table-cell paras)

    # Find body-children range that overlaps [comp_start, comp_end]
    print("\n  --- Body-level structure (paragraphs + tables) ---")
    cur_para = 0
    in_range = False
    for bi, child in enumerate(body_children):
        if child.tag == qn("w:p"):
            is_comp = comp_start <= cur_para <= comp_end
            if is_comp:
                in_range = True
                si = _sect_info(paragraphs[cur_para])
                text = repr(paragraphs[cur_para].text.strip()[:60] or "(blank)")
                print(f"  P[{cur_para:5d}] {text}{si}")
            elif in_range:
                break   # past compartment end
            cur_para += 1
        elif child.tag == qn("w:tbl"):
            # Count table-cell paragraphs
            tc_paras = list(child.iter(qn("w:p")))
            tc_count = len(tc_paras)
            if tc_paras:
                first_text = "".join(
                    t.text or "" for t in tc_paras[0].iter(qn("w:t")))
            else:
                first_text = ""
            if in_range:
                print(f"  {'':7}[TABLE  {tc_count} cell-paras, "
                      f"first: {repr(first_text[:40])}]")
            cur_para += tc_count


def _sect_info(para):
    pPr = para._element.pPr
    if pPr is None:
        return ""
    sp = pPr.find(qn("w:sectPr"))
    if sp is None:
        return ""
    te = sp.find(qn("w:type"))
    val = te.get(qn("w:val")) if te is not None else "nextPage(default)"
    cols_elem = sp.find(qn("w:cols"))
    num = cols_elem.get(qn("w:num")) if cols_elem is not None else "1"
    return f"  [sectPr type={val} cols={num}]"


def _para_line(i, para, markers=None):
    text = repr(para.text.strip()[:70] or "(blank)")
    tbl  = " [TABLE_CELL]" if _is_in_table(para) else ""
    sect = _sect_info(para)
    tag  = f"  <<<{markers}>>>" if markers else ""
    return f"  {i:5d}: {text}{tbl}{sect}{tag}"


def dump_compartment(doc, paragraphs, clauses, mapping, comp):
    print(f"\n{'='*70}")
    print(f"  {comp['name']}  (paras {comp['start']}–{comp['end']})")
    print(f"{'='*70}")

    # Match compartment in mapping (case-insensitive substring)
    short = comp["short_name"].lower()
    comp_key = next(
        (k for k in mapping if k.lower() == short or
         short in k.lower() or k.lower() in short),
        None,
    )
    clause_ids = mapping.get(comp_key, []) if comp_key else []

    if not clause_ids:
        print("  (no clauses mapped to this compartment)")

    for clause_id in clause_ids:
        if clause_id not in clauses:
            print(f"\n  Clause {clause_id}: NOT DEFINED in clauses sheet")
            continue
        clause = clauses[clause_id]
        anchor_idx = find_anchor(
            paragraphs, clause["anchor"], comp["start"], comp["end"])
        if anchor_idx is None:
            print(f"\n  Clause {clause_id}: anchor '{clause['anchor']}' NOT FOUND")
            continue

        insert_idx = find_insert_idx(
            paragraphs, anchor_idx, comp["end"], clause["position"])

        print(f"\n  Clause {clause_id} [{clause['type']}]"
              f"  anchor='{clause['anchor']}'  position={clause['position']}")
        print(f"  anchor_idx={anchor_idx}  →  insert_idx={insert_idx}")
        print()

        # Print context: 5 paras before anchor, anchor → insert+5
        start = max(comp["start"], anchor_idx - 5)
        end   = min(comp["end"],   insert_idx + 5)
        for i in range(start, end + 1):
            markers = ""
            if i == anchor_idx:
                markers = "ANCHOR"
            elif i == insert_idx:
                markers = "INSERT"
            print(_para_line(i, paragraphs[i], markers or None))

    # Always list every sectPr in the compartment
    print(f"\n  --- All sectPr paragraphs in this compartment ---")
    found = False
    for i in range(comp["start"], comp["end"] + 1):
        si = _sect_info(paragraphs[i])
        if si:
            found = True
            # Print 3 lines of context around it
            for j in range(max(comp["start"], i - 2), min(comp["end"] + 1, i + 3)):
                m = "sectPr" if j == i else None
                print(_para_line(j, paragraphs[j], m))
            print()
    if not found:
        print("  (none found)")

    # Show body-level structure (paragraphs + tables)
    dump_body_structure(doc, paragraphs, comp["start"], comp["end"])


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    docx_path  = sys.argv[1]
    xlsx_path  = sys.argv[2]
    comp_filter = sys.argv[3].lower() if len(sys.argv) > 3 else None

    doc = Document(docx_path)
    paragraphs = doc.paragraphs
    clauses, mapping = load_mapping(xlsx_path)
    compartments = find_compartments(doc)

    if comp_filter is None:
        print(f"Found {len(compartments)} compartments:\n")
        for c in compartments:
            short = c["short_name"]
            print(f"  {c['start']:5d}–{c['end']:5d}  {c['name']}")
        print("\nRe-run with a compartment name fragment to inspect it.")
        return

    matches = [c for c in compartments if comp_filter in c["name"].lower()]
    if not matches:
        print(f"No compartment matching '{comp_filter}'")
        return

    for comp in matches:
        dump_compartment(doc, paragraphs, clauses, mapping, comp)


if __name__ == "__main__":
    main()
