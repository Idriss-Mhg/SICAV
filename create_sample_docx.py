"""
Generates data/prospectus_sample.docx — a minimal fake prospectus
with 2 compartments structured like the real document.
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "data/prospectus_sample.docx"


def add(doc, text, style="Normal", bold=False, space_before=False):
    if space_before:
        doc.add_paragraph("", style="Normal")
    p = doc.add_paragraph(text, style=style)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def add_compartment(doc, number, name, currency, ref_currency, risk_method, investor_horizon):
    # --- Compartment title ---
    add(doc, f"SUPPLEMENT {number}. CPR Invest \u2013 {name}", style="Normal", bold=True)
    add(doc, "")

    add(doc, (
        f"The Compartment seeks to provide long-term capital growth by investing primarily "
        f"in equities and equity-related securities listed on recognised markets worldwide."
    ))
    add(doc, (
        "The Compartment may hold up to 15% of its assets in one or more of the following:"
    ))
    add(doc, "Ancillary liquid assets;")
    add(doc, "Financial derivative instruments, which may be used only for hedging purposes.")
    add(doc, "")

    add(doc, f"Launch date: 01/01/2015")
    add(doc, "")
    add(doc, "Term: Undetermined duration")
    add(doc, "")

    add(doc, f"Reference Currency: {ref_currency}")
    add(doc, "")

    # Investment strategy
    add(doc, "Investment strategy and policy of the Compartment", bold=True)
    add(doc, (
        f"The Compartment is actively managed and invests across a diversified portfolio "
        f"of equities. The management team uses both quantitative and qualitative analysis "
        f"to select securities."
    ))
    add(doc, "")

    add(doc, "Assets used by the Compartment", bold=True)
    add(doc, "Equities:")
    add(doc, (
        "The Compartment may invest up to 100% of its assets in equities and equity-related "
        "instruments listed on regulated markets."
    ))
    add(doc, "")
    add(doc, "Debt securities and money-market instruments:")
    add(doc, (
        "The Compartment may hold up to 20% of its assets in short-term money-market "
        "instruments for liquidity management purposes."
    ))
    add(doc, "")

    add(doc, "Derivatives in general used by the Compartment", bold=True)
    add(doc, (
        "The Compartment may use financial derivative instruments for hedging purposes only. "
        "The use of derivatives will not exceed a total commitment of one time the assets."
    ))
    add(doc, "")

    add(doc, "Total Return Swaps:")
    add(doc, "As an indication, total return swaps represent approximately 0% of net assets.")
    add(doc, "")

    add(doc, "Other transactions used by the Compartment", bold=True)
    add(doc, "Term deposits:")
    add(doc, "The Compartment may make term deposits with credit institutions for up to 12 months.")
    add(doc, "")
    add(doc, "Cash borrowings:")
    add(doc, "The Compartment may borrow up to 10% of its net assets for liquidity purposes.")
    add(doc, "")

    add(doc, "Transactions involving temporary acquisitions and/or disposals of securities:")
    add(doc, "Kinds of transaction used:")
    add(doc, "repo and reverse repo agreements;")
    add(doc, "lending and borrowing of securities.")
    add(doc, "")

    # Risk Management
    add(doc, "Risk Management:", bold=True)
    add(doc, f"The method used to calculate overall exposure of the Compartment is the {risk_method}.")
    add(doc, "")

    add(doc, "Interaction between the Compartment and the Master fund:")
    add(doc, "")
    add(doc, "Information flow between the Compartment and the Master Fund")
    add(doc, (
        "The Compartment and the Master Fund being both managed by the Management Company, "
        "they share all relevant information on a daily basis."
    ))
    add(doc, "")
    add(doc, "Information flow between the Depositary and the Master Fund's depositary bank")
    add(doc, "The Depositary and the Master Fund's depositary bank have entered into an information sharing agreement.")
    add(doc, "")

    # Profile
    add(doc, "Profile of typical investor in the Compartment:")
    add(doc, f"All investors who can afford to immobilize their capital for at least {investor_horizon} years;")
    add(doc, "accept to bear the risk of capital loss.")
    add(doc, "")

    add(doc, "Compartment's Benchmark: No")
    add(doc, "")

    add(doc, "Compartment's Performance Indicator:")
    add(doc, "The Compartment uses no formal benchmark. Performance is assessed on an absolute return basis.")
    add(doc, "")

    # Risk factors
    add(doc, "Compartment's Main and Specific Risk Factors:", bold=True)
    add(doc, "Main Risks")
    add(doc, "Capital Loss risk")
    add(doc, "Equity and Market risks")
    add(doc, "Liquidity risk")
    add(doc, "")
    add(doc, "Other risks")
    add(doc, "Counterparty risk")
    add(doc, "Credit risk")
    add(doc, "Exchange rate risk")
    add(doc, "Sustainable Investment Risk")
    add(doc, "")

    # Conflicts of interest
    add(doc, "Conflicts of interest")
    add(doc, (
        "The Management Company has implemented a conflicts of interest policy. "
        "Potential conflicts are managed through information barriers and escalation procedures."
    ))
    add(doc, "")

    # Shares characteristics
    add(doc, "Shares Characteristics / Subscription and Redemption conditions:")
    add(doc, f"Business Day: A full business day on which banks are open in Luxembourg and {currency} markets.")
    add(doc, "Valuation Day: Every Business Day.")
    add(doc, "Cut-off Time: 9.00 a.m. on the relevant Valuation Day.")
    add(doc, "Subscription and Redemption Settlement Day: 2 Business Days after the relevant Valuation Day.")
    add(doc, "")
    add(doc, "Main Share Classes")
    add(doc, "")

    # SFDR placeholder
    add(doc, f"SFDR Annex {number} - Pre-contractual disclosure for the financial products referred to in Article 8.")
    add(doc, "")
    add(doc, "")


def main():
    doc = Document()

    # --- Main body (minimal) ---
    add(doc, "PROSPECTUS", style="Heading 1")
    add(doc, "CPR Invest", style="Heading 2")
    add(doc, (
        "This Prospectus relates to CPR Invest (the \"Company\"), a Luxembourg société "
        "d'investissement à capital variable (SICAV) with multiple compartments."
    ))
    add(doc, "")

    add(doc, "Principal features", style="Heading 1")
    add(doc, "The Company offers a range of compartments with different investment objectives.")
    add(doc, "")

    add(doc, "Risk warnings", style="Heading 1")
    add(doc, "6.1. Introduction", style="Heading 2")
    add(doc, "Investment in the Company involves risk. Please read this section carefully.")
    add(doc, "")
    add(doc, "6.2. General risks", style="Heading 2")
    add(doc, "Market risk, liquidity risk, and counterparty risk apply to all compartments.")
    add(doc, "")

    # --- Appendices ---
    add(doc, "APPENDICES TO THE PROSPECTUS \u2013 COMPARTMENTS", style="Heading 1")
    add(doc, "")

    # --- Compartment 1 ---
    add_compartment(
        doc,
        number=1,
        name="Silver Age",
        currency="EUR",
        ref_currency="Euro (EUR)",
        risk_method="commitment approach",
        investor_horizon=5,
    )

    # --- Compartment 2 ---
    add_compartment(
        doc,
        number=2,
        name="Reactive",
        currency="EUR",
        ref_currency="Euro (EUR)",
        risk_method="value-at-risk (VaR) approach",
        investor_horizon=4,
    )

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
